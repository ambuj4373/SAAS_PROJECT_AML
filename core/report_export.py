"""
core/report_export.py
Professional PDF & DOCX report generator for Know Your Charity & Company UK.

Fonts: Times New Roman (DOCX) / Times-Roman (PDF built-in)
Body: 12 pt  |  Headings: 14 pt bold
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

# ── PDF (reportlab) ──────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, KeepTogether,
)

# ── DOCX (python-docx) ──────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

# ═════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═════════════════════════════════════════════════════════════════════════
_FONT_BODY = "Times-Roman"
_FONT_BOLD = "Times-Bold"
_FONT_ITALIC = "Times-Italic"
_FONT_DOCX = "Times New Roman"

_BODY_SIZE = 12
_HEADING1_SIZE = 18
_HEADING2_SIZE = 14
_HEADING3_SIZE = 13
_SMALL_SIZE = 10
_CAPTION_SIZE = 9

_PAGE_MARGIN = 2.5 * cm

_COLOR_PRIMARY = colors.HexColor("#1a365d")
_COLOR_ACCENT = colors.HexColor("#2b6cb0")
_COLOR_LIGHT_BG = colors.HexColor("#f7fafc")
_COLOR_BORDER = colors.HexColor("#cbd5e0")
_COLOR_RED = colors.HexColor("#c53030")
_COLOR_GREEN = colors.HexColor("#276749")
_COLOR_AMBER = colors.HexColor("#b7791f")

_DOCX_PRIMARY = RGBColor(0x1a, 0x36, 0x5d)
_DOCX_ACCENT = RGBColor(0x2b, 0x6c, 0xb0)
_DOCX_RED = RGBColor(0xc5, 0x30, 0x30)
_DOCX_GREEN = RGBColor(0x27, 0x67, 0x49)
_DOCX_AMBER = RGBColor(0xb7, 0x79, 0x1f)
_DOCX_GREY = RGBColor(0x71, 0x80, 0x96)


# ═════════════════════════════════════════════════════════════════════════
# HELPERS
# ═════════════════════════════════════════════════════════════════════════

def _fmt_money(val: Any) -> str:
    """Format a monetary value with £ sign and comma separators."""
    try:
        v = float(val or 0)
        if v >= 1_000_000:
            return f"£{v / 1_000_000:,.2f}M"
        if v >= 1_000:
            return f"£{v / 1_000:,.1f}K"
        return f"£{v:,.0f}"
    except (ValueError, TypeError):
        return "£0"


def _risk_color_pdf(level: str):
    """Return reportlab color for a risk level string."""
    up = (level or "").upper().replace("-", " ")
    if "HIGH" in up:
        return _COLOR_RED
    if "MEDIUM" in up or "AMBER" in up:
        return _COLOR_AMBER
    if "LOW" in up:
        return _COLOR_GREEN
    return colors.grey


def _risk_color_docx(level: str) -> RGBColor:
    """Return DOCX RGBColor for a risk level string."""
    up = (level or "").upper().replace("-", " ")
    if "HIGH" in up:
        return _DOCX_RED
    if "MEDIUM" in up or "AMBER" in up:
        return _DOCX_AMBER
    if "LOW" in up:
        return _DOCX_GREEN
    return _DOCX_GREY


def _clean_md(text: str) -> str:
    """Strip markdown fences, HTML tags, and emoji for clean PDF/DOCX text."""
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    # Convert markdown links [text](url) → text
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)
    return text.strip()


def _parse_md_table_row(line: str) -> List[str]:
    """Parse a markdown table row '| a | b | c |' into ['a', 'b', 'c']."""
    # Strip leading/trailing pipe and split
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _is_md_table_separator(line: str) -> bool:
    """Check if line is a markdown table separator like |---|---|."""
    return bool(re.match(r"^\s*\|[\s:*-]+(\|[\s:*-]+)+\|?\s*$", line))


def _md_to_paragraphs(md_text: str) -> List[Dict[str, Any]]:
    """
    Parse markdown text into a list of paragraph dicts:
      {"type": "h1"|"h2"|"h3"|"bullet"|"text"|"table", "content": str|dict}
    Handles: ## headings, **bold**, - bullets, plain paragraphs, | tables.
    """
    paragraphs = []
    if not md_text:
        return paragraphs

    lines = md_text.split("\n")
    buffer = []
    table_rows: List[List[str]] = []

    def flush():
        if buffer:
            joined = " ".join(buffer).strip()
            if joined:
                paragraphs.append({"type": "text", "content": joined})
            buffer.clear()

    def flush_table():
        if table_rows:
            # First row is headers, rest are data
            headers = table_rows[0] if table_rows else []
            data = table_rows[1:] if len(table_rows) > 1 else []
            paragraphs.append({
                "type": "table",
                "content": {"headers": headers, "rows": data},
            })
            table_rows.clear()

    for line in lines:
        stripped = line.strip()

        # Check if this is a table line (starts with |)
        if stripped.startswith("|") and "|" in stripped[1:]:
            # If we were accumulating text, flush it first
            flush()

            # Skip separator rows (|---|---|)
            if _is_md_table_separator(stripped):
                continue

            cells = _parse_md_table_row(stripped)
            # Filter out empty rows from bad formatting
            if any(c.strip() for c in cells):
                table_rows.append(cells)
            continue

        # If we were building a table and hit a non-table line, flush the table
        if table_rows:
            flush_table()

        # Empty line → flush paragraph buffer
        if not stripped:
            flush()
            continue

        # Headings
        if stripped.startswith("####"):
            flush()
            paragraphs.append({"type": "h3", "content": stripped.lstrip("#").strip()})
        elif stripped.startswith("###"):
            flush()
            paragraphs.append({"type": "h3", "content": stripped.lstrip("#").strip()})
        elif stripped.startswith("##"):
            flush()
            paragraphs.append({"type": "h2", "content": stripped.lstrip("#").strip()})
        elif stripped.startswith("#"):
            flush()
            paragraphs.append({"type": "h1", "content": stripped.lstrip("#").strip()})
        # Horizontal rules
        elif re.match(r"^[-*_]{3,}$", stripped):
            flush()
            paragraphs.append({"type": "hr", "content": ""})
        # Bullets
        elif re.match(r"^[-*•]\s+", stripped):
            flush()
            content = re.sub(r"^[-*•]\s+", "", stripped)
            paragraphs.append({"type": "bullet", "content": content})
        # Numbered list
        elif re.match(r"^\d+[.)]\s+", stripped):
            flush()
            paragraphs.append({"type": "bullet", "content": stripped})
        # Blockquote
        elif stripped.startswith(">"):
            flush()
            content = stripped.lstrip(">").strip()
            if content:
                paragraphs.append({"type": "text", "content": content})
        else:
            buffer.append(stripped)

    flush()
    flush_table()
    return paragraphs


def _md_inline_to_rl(text: str) -> str:
    """Convert **bold** and *italic* to reportlab XML tags, escape XML chars."""
    # Strip remaining markdown links [text](url) → text
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)
    # Escape XML-special characters BEFORE adding our own tags
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;").replace(">", "&gt;")
    # Bold
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)
    # Strip remaining markdown artifacts
    text = text.replace("***", "")
    return text


# ═════════════════════════════════════════════════════════════════════════
# PDF STYLES
# ═════════════════════════════════════════════════════════════════════════

def _build_pdf_styles():
    """Build custom paragraph styles for the PDF report."""
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "CoverTitle",
        fontName=_FONT_BOLD,
        fontSize=22,
        leading=28,
        alignment=TA_CENTER,
        textColor=_COLOR_PRIMARY,
        spaceAfter=6 * mm,
    ))
    styles.add(ParagraphStyle(
        "CoverSubtitle",
        fontName=_FONT_BODY,
        fontSize=14,
        leading=20,
        alignment=TA_CENTER,
        textColor=_COLOR_ACCENT,
        spaceAfter=4 * mm,
    ))
    styles.add(ParagraphStyle(
        "CoverMeta",
        fontName=_FONT_BODY,
        fontSize=11,
        leading=16,
        alignment=TA_CENTER,
        textColor=colors.grey,
        spaceAfter=2 * mm,
    ))
    styles.add(ParagraphStyle(
        "SectionHeading",
        fontName=_FONT_BOLD,
        fontSize=_HEADING2_SIZE,
        leading=_HEADING2_SIZE + 4,
        textColor=_COLOR_PRIMARY,
        spaceBefore=10 * mm,
        spaceAfter=4 * mm,
        borderWidth=0,
        borderPadding=0,
    ))
    styles.add(ParagraphStyle(
        "SubHeading",
        fontName=_FONT_BOLD,
        fontSize=_HEADING3_SIZE,
        leading=_HEADING3_SIZE + 3,
        textColor=_COLOR_ACCENT,
        spaceBefore=6 * mm,
        spaceAfter=2 * mm,
    ))
    styles.add(ParagraphStyle(
        "BodyText14",
        fontName=_FONT_BODY,
        fontSize=_BODY_SIZE,
        leading=_BODY_SIZE + 6,
        alignment=TA_JUSTIFY,
        spaceAfter=3 * mm,
    ))
    styles.add(ParagraphStyle(
        "Bullet14",
        fontName=_FONT_BODY,
        fontSize=_BODY_SIZE,
        leading=_BODY_SIZE + 5,
        leftIndent=12 * mm,
        bulletIndent=6 * mm,
        spaceBefore=1 * mm,
        spaceAfter=1 * mm,
    ))
    styles.add(ParagraphStyle(
        "SmallText",
        fontName=_FONT_BODY,
        fontSize=_SMALL_SIZE,
        leading=_SMALL_SIZE + 4,
        textColor=colors.grey,
        spaceAfter=2 * mm,
    ))
    styles.add(ParagraphStyle(
        "RiskHigh",
        fontName=_FONT_BOLD,
        fontSize=_BODY_SIZE,
        textColor=_COLOR_RED,
    ))
    styles.add(ParagraphStyle(
        "RiskLow",
        fontName=_FONT_BOLD,
        fontSize=_BODY_SIZE,
        textColor=_COLOR_GREEN,
    ))
    styles.add(ParagraphStyle(
        "Disclaimer",
        fontName=_FONT_ITALIC,
        fontSize=_CAPTION_SIZE,
        leading=_CAPTION_SIZE + 4,
        textColor=colors.grey,
        alignment=TA_JUSTIFY,
        spaceBefore=8 * mm,
    ))
    return styles


# ═════════════════════════════════════════════════════════════════════════
# PDF GENERATION — CHARITY
# ═════════════════════════════════════════════════════════════════════════

def generate_charity_pdf(dp: Dict[str, Any]) -> bytes:
    """
    Generate a professional PDF report for a charity.
    Args:
        dp: The _display dict from st.session_state
    Returns:
        PDF file content as bytes.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=_PAGE_MARGIN,
        rightMargin=_PAGE_MARGIN,
        topMargin=_PAGE_MARGIN,
        bottomMargin=_PAGE_MARGIN,
        title="Know Your Charity UK — Due Diligence Report",
    )
    styles = _build_pdf_styles()
    story: List = []

    # ── Extract data ─────────────────────────────────────────────────
    entity_name = dp.get("entity_name", "Unknown Charity")
    charity_num = dp.get("charity_num", "")
    charity_data = dp.get("charity_data", {})
    full_report = dp.get("full_report", "")
    trustees = dp.get("trustees", [])
    financial_history = dp.get("financial_history", [])
    financial_anomalies = dp.get("financial_anomalies", {})
    countries = dp.get("countries", [])
    high_risk_countries = dp.get("high_risk_countries", [])
    hrcob_core_controls = dp.get("hrcob_core_controls", {})
    adverse_org = dp.get("adverse_org", {})
    adverse_trustees = dp.get("adverse_trustees", {})
    cost_info = dp.get("cost_info", {})
    v3_risk = dp.get("v3_risk_score", {})
    report_date = datetime.now().strftime("%d %B %Y")

    # Overall risk from LLM report
    _risk_match = re.search(
        r"(?:Overall\s+Risk\s+Rating)[:\s*#]*\*{0,2}\s*"
        r"(LOW|MEDIUM|HIGH|VERY\s*HIGH|MEDIUM[\s-]*HIGH|MEDIUM[\s-]*LOW)",
        full_report, re.IGNORECASE,
    )
    overall_risk = _risk_match.group(1).strip().upper() if _risk_match else "UNRATED"
    hrcob_stat = hrcob_core_controls.get("hrcob_status", "Unknown")
    inc = charity_data.get("latest_income") or 0
    exp = charity_data.get("latest_expenditure") or 0

    # ── COVER PAGE ───────────────────────────────────────────────────
    story.append(Spacer(1, 40 * mm))
    story.append(HRFlowable(
        width="80%", thickness=2, color=_COLOR_PRIMARY,
        spaceAfter=6 * mm, spaceBefore=0,
    ))
    story.append(Paragraph("Know Your Charity UK", styles["CoverTitle"]))
    story.append(Paragraph("In-Depth Due Diligence Report", styles["CoverSubtitle"]))
    story.append(HRFlowable(
        width="80%", thickness=2, color=_COLOR_PRIMARY,
        spaceAfter=10 * mm, spaceBefore=0,
    ))
    story.append(Spacer(1, 10 * mm))
    story.append(Paragraph(entity_name, styles["CoverTitle"]))
    story.append(Spacer(1, 6 * mm))
    _cover_meta = [
        f"Charity Commission No: {charity_num}",
        f"Report Date: {report_date}",
        f"Registration Status: {charity_data.get('reg_status', 'Unknown')}",
    ]
    for m in _cover_meta:
        story.append(Paragraph(m, styles["CoverMeta"]))
    story.append(Spacer(1, 15 * mm))
    story.append(Paragraph(
        f"Overall Risk Assessment: {overall_risk}",
        ParagraphStyle(
            "CoverRisk", parent=styles["CoverSubtitle"],
            textColor=_risk_color_pdf(overall_risk),
            fontName=_FONT_BOLD, fontSize=16,
        ),
    ))
    story.append(Spacer(1, 30 * mm))
    story.append(Paragraph(
        "CONFIDENTIAL — For authorised recipients only",
        styles["SmallText"],
    ))
    story.append(PageBreak())

    # ── TABLE OF CONTENTS (lightweight) ──────────────────────────────
    story.append(Paragraph("Contents", styles["SectionHeading"]))
    toc_items = [
        "1. Executive Summary",
        "2. Financial Overview",
        "3. Governance & Controls",
        "4. Geography Risk Assessment",
        "5. Screening Results",
        "6. AI Analyst Report",
        "7. Disclaimer",
    ]
    for item in toc_items:
        story.append(Paragraph(item, styles["BodyText14"]))
    story.append(PageBreak())

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────
    story.append(Paragraph("1. Executive Summary", styles["SectionHeading"]))

    # Overall risk
    story.append(Paragraph(
        f"<b>Overall Risk Rating:</b> {overall_risk}",
        ParagraphStyle("_esr", parent=styles["BodyText14"],
                        textColor=_risk_color_pdf(overall_risk)),
    ))

    # HRCOB status
    story.append(Paragraph(
        f"<b>HRCOB Core Controls Status:</b> {hrcob_stat}",
        styles["BodyText14"],
    ))

    # Key metrics table
    _exec_data = [
        ["Metric", "Value"],
        ["Latest Income", _fmt_money(inc)],
        ["Latest Expenditure", _fmt_money(exp)],
        ["Surplus / (Deficit)", _fmt_money(inc - exp)],
        ["Number of Trustees", str(len(trustees))],
        ["Countries of Operation", str(len(countries))],
        ["High-Risk Jurisdictions", str(len(high_risk_countries))],
        ["Financial Anomalies", str(financial_anomalies.get("anomaly_count", 0) if financial_anomalies else 0)],
    ]
    t = Table(_exec_data, colWidths=[200, 200])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
        ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
        ("FONTSIZE", (0, 0), (-1, -1), _BODY_SIZE),
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 1), (-1, -1), _COLOR_LIGHT_BG),
        ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(Spacer(1, 4 * mm))
    story.append(t)

    # Key findings
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("<b>Key Findings:</b>", styles["BodyText14"]))
    _findings = []
    _findings.append(f"HRCOB Core Controls: {hrcob_stat}")
    if high_risk_countries:
        _findings.append(f"{len(high_risk_countries)} high-risk jurisdiction(s): {', '.join(high_risk_countries[:5])}")
    else:
        _findings.append("No high-risk jurisdictions identified")
    _n_adverse = sum(_count_adverse(v) for v in adverse_trustees.values()) + _count_adverse(adverse_org)
    if _n_adverse > 0:
        _findings.append(f"{_n_adverse} verified adverse media hit(s) — review required")
    else:
        _findings.append("No verified adverse media hits")
    if inc > 0:
        _spend_pct = exp / inc * 100
        _findings.append(
            f"Income {_fmt_money(inc)}, Spend-to-Income {_spend_pct:.0f}%, "
            f"{'Surplus' if inc >= exp else 'Deficit'} {_fmt_money(abs(inc - exp))}"
        )
    for f in _findings:
        story.append(Paragraph(f"• {f}", styles["Bullet14"]))

    # ── 2. FINANCIAL OVERVIEW ────────────────────────────────────────
    story.append(Paragraph("2. Financial Overview", styles["SectionHeading"]))

    # Income breakdown
    _inc_sources = {
        "Donations & Legacies": charity_data.get("inc_donations", 0) or 0,
        "Charitable Activities": charity_data.get("inc_charitable", 0) or 0,
        "Trading": charity_data.get("inc_trading", 0) or 0,
        "Investments": charity_data.get("inc_investments", 0) or 0,
        "Other": charity_data.get("inc_other", 0) or 0,
    }
    _inc_rows = [["Source", "Amount"]]
    for src, amt in _inc_sources.items():
        _inc_rows.append([src, _fmt_money(amt)])
    _inc_rows.append(["Total Income", _fmt_money(inc)])

    t2 = Table(_inc_rows, colWidths=[240, 160])
    t2.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
        ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
        ("FONTNAME", (0, -1), (-1, -1), _FONT_BOLD),
        ("FONTSIZE", (0, 0), (-1, -1), _BODY_SIZE),
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#edf2f7")),
        ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(Paragraph("<b>Income Breakdown (Latest Year)</b>", styles["SubHeading"]))
    story.append(t2)

    # Financial history
    if financial_history:
        story.append(Paragraph("<b>Financial History</b>", styles["SubHeading"]))
        _fh_rows = [["Year End", "Income", "Expenditure", "Surplus/Deficit"]]
        for fh in financial_history[:6]:
            _yr = str(fh.get("year_end", ""))[:10] if fh.get("year_end") else "—"
            _fi = fh.get("income") or 0
            _fe = fh.get("expenditure") or 0
            _fh_rows.append([_yr, _fmt_money(_fi), _fmt_money(_fe), _fmt_money(_fi - _fe)])

        t3 = Table(_fh_rows, colWidths=[100, 110, 110, 110])
        t3.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
            ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
            ("FONTSIZE", (0, 0), (-1, -1), 11),
            ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
        ]))
        story.append(t3)

    # Anomalies
    _anomaly_ct = financial_anomalies.get("anomaly_count", 0) if financial_anomalies else 0
    if _anomaly_ct > 0:
        story.append(Paragraph(
            f"<b>Financial Anomalies Detected:</b> {_anomaly_ct}",
            styles["BodyText14"],
        ))
        for det in (financial_anomalies.get("details") or [])[:5]:
            story.append(Paragraph(f"• {det}", styles["Bullet14"]))

    # ── 3. GOVERNANCE & CONTROLS ─────────────────────────────────────
    story.append(Paragraph("3. Governance & Controls", styles["SectionHeading"]))
    story.append(Paragraph(
        f"<b>HRCOB Core Controls Assessment:</b> {hrcob_stat}",
        styles["BodyText14"],
    ))

    # Core controls table
    _ctrl_areas = ["safeguarding", "financial_crime", "risk_management"]
    _ctrl_rows = [["Control Area", "Status", "Evidence"]]
    for area in _ctrl_areas:
        ctrl = hrcob_core_controls.get(area, {})
        _status = ctrl.get("status", "not_located").replace("_", " ").title()
        _evidence = ctrl.get("evidence", "—")
        if len(_evidence) > 120:
            _evidence = _evidence[:117] + "..."
        _ctrl_rows.append([area.replace("_", " ").title(), _status, _evidence])

    t4 = Table(_ctrl_rows, colWidths=[120, 90, 230])
    t4.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
        ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(t4)

    # Trustees summary
    if trustees:
        story.append(Paragraph("<b>Board of Trustees</b>", styles["SubHeading"]))
        _t_rows = [["Name", "Role"]]
        for tr in trustees[:15]:
            _tname = tr.get("name", "—") if isinstance(tr, dict) else str(tr)
            _trole = tr.get("role", "Trustee") if isinstance(tr, dict) else "Trustee"
            _t_rows.append([_tname, _trole])
        if len(trustees) > 15:
            _t_rows.append([f"... and {len(trustees) - 15} more", ""])

        t5 = Table(_t_rows, colWidths=[260, 160])
        t5.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
            ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
            ("FONTSIZE", (0, 0), (-1, -1), 11),
            ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
        ]))
        story.append(t5)

    # ── 4. GEOGRAPHY ─────────────────────────────────────────────────
    story.append(Paragraph("4. Geography Risk Assessment", styles["SectionHeading"]))
    _country_names = [c.get("country", str(c)) if isinstance(c, dict) else str(c) for c in countries]
    if _country_names:
        story.append(Paragraph(
            f"<b>Countries of Operation:</b> {', '.join(_country_names[:20])}",
            styles["BodyText14"],
        ))
    else:
        story.append(Paragraph("No international operations identified.", styles["BodyText14"]))

    if high_risk_countries:
        story.append(Paragraph(
            f"<b>High-Risk Jurisdictions ({len(high_risk_countries)}):</b>",
            styles["BodyText14"],
        ))
        for c in high_risk_countries:
            story.append(Paragraph(f"• {c}", styles["Bullet14"]))
    else:
        story.append(Paragraph(
            "No high-risk jurisdictions identified.",
            styles["BodyText14"],
        ))

    # ── 5. SCREENING RESULTS ─────────────────────────────────────────
    story.append(Paragraph("5. Screening Results", styles["SectionHeading"]))
    story.append(Paragraph("<b>Adverse Media Screening</b>", styles["SubHeading"]))
    if _n_adverse > 0:
        story.append(Paragraph(
            f"{_n_adverse} verified adverse media hit(s) detected. Manual review recommended.",
            ParagraphStyle("_adv", parent=styles["BodyText14"], textColor=_COLOR_RED),
        ))
    else:
        story.append(Paragraph(
            "No verified adverse media hits detected for the organisation or its trustees.",
            styles["BodyText14"],
        ))

    # FATF screening summary
    fatf_org = dp.get("fatf_org_screen", {})
    if fatf_org:
        _fatf_result = fatf_org.get("result", "No Match")
        story.append(Paragraph("<b>FATF Screening</b>", styles["SubHeading"]))
        story.append(Paragraph(f"Organisation FATF Result: {_fatf_result}", styles["BodyText14"]))

    # ── 6. AI ANALYST REPORT ─────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("6. AI Analyst Report", styles["SectionHeading"]))
    story.append(Paragraph(
        "The following analyst report was generated by AI based on structured data, "
        "Charity Commission filings, Companies House records, web searches, and "
        "document analysis.",
        styles["SmallText"],
    ))

    # Parse and render the markdown report
    _paras = _md_to_paragraphs(_clean_md(full_report))
    for p in _paras:
        if p["type"] == "h1":
            story.append(Paragraph(p["content"], styles["SectionHeading"]))
        elif p["type"] == "h2":
            story.append(Paragraph(p["content"], styles["SubHeading"]))
        elif p["type"] == "h3":
            story.append(Paragraph(
                f"<b>{p['content']}</b>",
                ParagraphStyle("_h3", parent=styles["BodyText14"], fontName=_FONT_BOLD),
            ))
        elif p["type"] == "bullet":
            story.append(Paragraph(
                f"• {_md_inline_to_rl(p['content'])}",
                styles["Bullet14"],
            ))
        elif p["type"] == "hr":
            story.append(HRFlowable(
                width="100%", thickness=0.5, color=_COLOR_BORDER,
                spaceAfter=3 * mm, spaceBefore=3 * mm,
            ))
        elif p["type"] == "table":
            _tbl_data = p["content"]
            _tbl_headers = _tbl_data.get("headers", [])
            _tbl_rows = _tbl_data.get("rows", [])
            if _tbl_headers:
                # Build ReportLab table
                _num_cols = len(_tbl_headers)
                _all_rows = [_tbl_headers] + _tbl_rows
                # Wrap cell text in Paragraphs for word-wrap
                _cell_style = ParagraphStyle(
                    "_tblcell", parent=styles["BodyText14"],
                    fontSize=9, leading=12, alignment=TA_LEFT, spaceAfter=0,
                )
                _cell_style_hdr = ParagraphStyle(
                    "_tblhdr", parent=_cell_style,
                    fontName=_FONT_BOLD, textColor=colors.white,
                )
                _rl_rows = []
                for ri, row in enumerate(_all_rows):
                    # Pad/trim row to match header column count
                    padded = (row + [""] * _num_cols)[:_num_cols]
                    if ri == 0:
                        _rl_rows.append([
                            Paragraph(_md_inline_to_rl(c), _cell_style_hdr) for c in padded
                        ])
                    else:
                        _rl_rows.append([
                            Paragraph(_md_inline_to_rl(c), _cell_style) for c in padded
                        ])
                # Calculate column widths (distribute available width)
                _avail = A4[0] - 2 * _PAGE_MARGIN
                _col_widths = [_avail / _num_cols] * _num_cols
                _tbl_obj = Table(_rl_rows, colWidths=_col_widths, repeatRows=1)
                _tbl_obj.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
                    ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
                    ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(Spacer(1, 2 * mm))
                story.append(_tbl_obj)
                story.append(Spacer(1, 2 * mm))
        else:
            story.append(Paragraph(_md_inline_to_rl(p["content"]), styles["BodyText14"]))

    # ── 7. DISCLAIMER ────────────────────────────────────────────────
    story.append(Spacer(1, 10 * mm))
    story.append(HRFlowable(width="100%", thickness=1, color=_COLOR_BORDER, spaceAfter=4 * mm))
    story.append(Paragraph("7. Disclaimer", styles["SectionHeading"]))
    story.append(Paragraph(
        "This report has been generated using automated data retrieval and AI-powered analysis. "
        "While every effort is made to ensure accuracy, the information contained herein is based on "
        "publicly available data sources including the Charity Commission for England and Wales, "
        "Companies House, FATF lists, and web-based searches. This report does not constitute legal, "
        "financial, or regulatory advice. Recipients should independently verify key findings before "
        "making decisions. The overall risk rating reflects a point-in-time assessment and may change "
        "as new information becomes available. AI-generated narrative sections may contain inaccuracies "
        "and should be reviewed by a qualified analyst.",
        styles["Disclaimer"],
    ))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(
        f"Generated by Know Your Charity UK on {report_date}. "
        "Built by Ambuj Shukla. Confidential.",
        styles["Disclaimer"],
    ))

    # ── Build PDF ────────────────────────────────────────────────────
    doc.build(story)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════
# PDF GENERATION — COMPANY
# ═════════════════════════════════════════════════════════════════════════

def generate_company_pdf(co_dp: Dict[str, Any]) -> bytes:
    """
    Generate a professional PDF report for a company.
    Args:
        co_dp: The _co_display dict from st.session_state
    Returns:
        PDF file content as bytes.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=_PAGE_MARGIN,
        rightMargin=_PAGE_MARGIN,
        topMargin=_PAGE_MARGIN,
        bottomMargin=_PAGE_MARGIN,
        title="Know Your Company UK — Due Diligence Report",
    )
    styles = _build_pdf_styles()
    story: List = []

    # ── Extract data ─────────────────────────────────────────────────
    co_data = co_dp.get("co_check_data", {})
    co_report = co_dp.get("co_report", "")
    co_cost = co_dp.get("co_cost_info", {})
    v3_risk = co_dp.get("v3_risk_score", {})
    report_date = datetime.now().strftime("%d %B %Y")

    co_name = co_data.get("company_name", "Unknown Company")
    co_num = co_data.get("company_number", "")
    co_profile = co_data.get("profile", {})
    co_risk = co_data.get("risk_matrix", {})
    co_age = co_data.get("company_age", {})
    co_directors = co_data.get("director_analysis", {})
    co_pscs = co_data.get("psc_analysis", {})
    co_xref = co_data.get("cross_reference", {})
    co_accounts = co_data.get("accounts_data", {})
    co_merchant = co_data.get("merchant_suitability", {})
    co_sic = co_data.get("sic_risk", {})
    co_virtual = co_data.get("virtual_office", {})
    co_ubo = co_data.get("ubo_chain", {})

    _co_overall = co_risk.get("overall_risk", "Unknown")
    _co_flags = co_risk.get("total_flags", 0)
    _co_score = co_risk.get("risk_score", "N/A")
    _co_hard_stops = co_risk.get("hard_stops", [])

    # ── COVER PAGE ───────────────────────────────────────────────────
    story.append(Spacer(1, 40 * mm))
    story.append(HRFlowable(
        width="80%", thickness=2, color=_COLOR_PRIMARY,
        spaceAfter=6 * mm,
    ))
    story.append(Paragraph("Know Your Company UK", styles["CoverTitle"]))
    story.append(Paragraph("Company Sense-Check Report", styles["CoverSubtitle"]))
    story.append(HRFlowable(
        width="80%", thickness=2, color=_COLOR_PRIMARY,
        spaceAfter=10 * mm,
    ))
    story.append(Spacer(1, 10 * mm))
    story.append(Paragraph(co_name, styles["CoverTitle"]))
    story.append(Spacer(1, 6 * mm))
    _cover_meta = [
        f"Companies House No: {co_num}",
        f"Report Date: {report_date}",
        f"Company Status: {co_profile.get('status', 'Unknown')}",
    ]
    for m in _cover_meta:
        story.append(Paragraph(m, styles["CoverMeta"]))
    story.append(Spacer(1, 15 * mm))
    story.append(Paragraph(
        f"Overall Risk: {_co_overall.upper()} — Score: {_co_score}/100",
        ParagraphStyle(
            "CoverRisk", parent=styles["CoverSubtitle"],
            textColor=_risk_color_pdf(_co_overall),
            fontName=_FONT_BOLD, fontSize=16,
        ),
    ))
    story.append(Spacer(1, 30 * mm))
    story.append(Paragraph(
        "CONFIDENTIAL — For authorised recipients only",
        styles["SmallText"],
    ))
    story.append(PageBreak())

    # ── TABLE OF CONTENTS ────────────────────────────────────────────
    story.append(Paragraph("Contents", styles["SectionHeading"]))
    toc_items = [
        "1. Executive Summary",
        "2. Company Profile",
        "3. Directors & Ownership",
        "4. Risk Matrix",
        "5. AI Analyst Report",
        "6. Disclaimer",
    ]
    for item in toc_items:
        story.append(Paragraph(item, styles["BodyText14"]))
    story.append(PageBreak())

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────
    story.append(Paragraph("1. Executive Summary", styles["SectionHeading"]))
    story.append(Paragraph(
        f"<b>Overall Risk:</b> {_co_overall.upper()}",
        ParagraphStyle("_cr", parent=styles["BodyText14"],
                        textColor=_risk_color_pdf(_co_overall)),
    ))
    story.append(Paragraph(
        f"<b>Risk Score:</b> {_co_score}/100 &nbsp; | &nbsp; "
        f"<b>Flags Detected:</b> {_co_flags}",
        styles["BodyText14"],
    ))

    if _co_hard_stops:
        story.append(Paragraph("<b>Hard Stops:</b>", styles["SubHeading"]))
        for hs in _co_hard_stops:
            story.append(Paragraph(
                f"• {hs}",
                ParagraphStyle("_chs", parent=styles["Bullet14"], textColor=_COLOR_RED),
            ))

    _exec_data = [
        ["Metric", "Value"],
        ["Company Status", co_profile.get("status", "—")],
        ["Company Type", co_profile.get("type", "—")],
        ["Company Age", f"{co_age.get('years', '?')} year(s)"],
        ["SIC Codes", ", ".join(co_profile.get("sic_codes", []) or ["—"])],
        ["Director Count", str(co_directors.get("director_count", 0))],
        ["UBO Layers Traced", str(co_ubo.get("layers_traced", 0))],
        ["Business Model", co_merchant.get("business_model", "—")],
        ["Website Credibility", co_xref.get("credibility_level", "Unknown")],
    ]
    t = Table(_exec_data, colWidths=[200, 230])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
        ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
        ("FONTSIZE", (0, 0), (-1, -1), _BODY_SIZE),
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 1), (-1, -1), _COLOR_LIGHT_BG),
        ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(Spacer(1, 4 * mm))
    story.append(t)

    # ── 2. COMPANY PROFILE ───────────────────────────────────────────
    story.append(Paragraph("2. Company Profile", styles["SectionHeading"]))
    _prof_items = [
        ("Legal Name", co_name),
        ("Company Number", co_num),
        ("Incorporated", co_profile.get("date_of_creation", "—")),
        ("Jurisdiction", co_profile.get("jurisdiction", "—")),
        ("Accounts Next Due", co_profile.get("accounts_next_due", "—")),
        ("Confirmation Next Due", co_profile.get("confirmation_next_due", "—")),
    ]
    _reg = co_profile.get("registered_office", {})
    if _reg:
        _addr = ", ".join(
            _reg.get(k, "") for k in
            ["address_line_1", "address_line_2", "locality", "region", "postal_code"]
            if _reg.get(k)
        )
        _prof_items.append(("Registered Office", _addr))

    _p_rows = [["Field", "Value"]] + [[k, v] for k, v in _prof_items]
    t2 = Table(_p_rows, colWidths=[160, 270])
    t2.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
        ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
    ]))
    story.append(t2)

    # SIC risk + virtual office
    if co_sic:
        story.append(Paragraph(
            f"<b>SIC Risk Level:</b> {co_sic.get('risk_level', 'Unknown')}",
            styles["BodyText14"],
        ))
    if co_virtual:
        story.append(Paragraph(
            f"<b>Virtual Office Indicator:</b> {co_virtual.get('indicator', 'Unknown')}",
            styles["BodyText14"],
        ))

    # ── 3. DIRECTORS & OWNERSHIP ─────────────────────────────────────
    story.append(Paragraph("3. Directors & Ownership", styles["SectionHeading"]))

    # Directors
    _dirs = co_directors.get("directors", [])
    if _dirs:
        _d_rows = [["Name", "Role", "Appointed", "Nationality"]]
        for d in _dirs[:12]:
            _d_rows.append([
                d.get("name", "—"),
                d.get("officer_role", "—"),
                str(d.get("appointed_on", "—")),
                d.get("nationality", "—"),
            ])
        if len(_dirs) > 12:
            _d_rows.append([f"... and {len(_dirs) - 12} more", "", "", ""])

        t3 = Table(_d_rows, colWidths=[150, 90, 90, 100])
        t3.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
            ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
        ]))
        story.append(Paragraph("<b>Active Directors</b>", styles["SubHeading"]))
        story.append(t3)

    # PSCs
    _pscs = co_pscs.get("pscs", [])
    if _pscs:
        _psc_rows = [["Name", "Nationality", "Ownership"]]
        for p in _pscs[:10]:
            _psc_rows.append([
                p.get("name", "—"),
                p.get("nationality", "—"),
                p.get("ownership_band", "—"),
            ])
        t4 = Table(_psc_rows, colWidths=[180, 120, 130])
        t4.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
            ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(Paragraph("<b>Persons with Significant Control</b>", styles["SubHeading"]))
        story.append(t4)

    # ── 4. RISK MATRIX ───────────────────────────────────────────────
    story.append(Paragraph("4. Risk Matrix", styles["SectionHeading"]))
    _cat_risks = co_risk.get("category_risks", {})
    if _cat_risks:
        _rm_rows = [["Category", "Risk Level"]]
        for cat, level in _cat_risks.items():
            _rm_rows.append([cat, level.title()])
        t5 = Table(_rm_rows, colWidths=[250, 180])
        t5.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
            ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
            ("FONTSIZE", (0, 0), (-1, -1), _BODY_SIZE),
            ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
        ]))
        story.append(t5)

    # ── 5. AI ANALYST REPORT ─────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("5. AI Analyst Report", styles["SectionHeading"]))
    story.append(Paragraph(
        "The following analyst report was generated by AI based on Companies House "
        "data, web searches, and cross-referencing analysis.",
        styles["SmallText"],
    ))

    _paras = _md_to_paragraphs(_clean_md(co_report))
    for p in _paras:
        if p["type"] == "h1":
            story.append(Paragraph(p["content"], styles["SectionHeading"]))
        elif p["type"] == "h2":
            story.append(Paragraph(p["content"], styles["SubHeading"]))
        elif p["type"] == "h3":
            story.append(Paragraph(
                f"<b>{p['content']}</b>",
                ParagraphStyle("_h3c", parent=styles["BodyText14"], fontName=_FONT_BOLD),
            ))
        elif p["type"] == "bullet":
            story.append(Paragraph(
                f"• {_md_inline_to_rl(p['content'])}",
                styles["Bullet14"],
            ))
        elif p["type"] == "hr":
            story.append(HRFlowable(
                width="100%", thickness=0.5, color=_COLOR_BORDER,
                spaceAfter=3 * mm, spaceBefore=3 * mm,
            ))
        elif p["type"] == "table":
            _tbl_data = p["content"]
            _tbl_headers = _tbl_data.get("headers", [])
            _tbl_rows = _tbl_data.get("rows", [])
            if _tbl_headers:
                _num_cols = len(_tbl_headers)
                _all_rows = [_tbl_headers] + _tbl_rows
                _cell_style = ParagraphStyle(
                    "_ctblcell", parent=styles["BodyText14"],
                    fontSize=9, leading=12, alignment=TA_LEFT, spaceAfter=0,
                )
                _cell_style_hdr = ParagraphStyle(
                    "_ctblhdr", parent=_cell_style,
                    fontName=_FONT_BOLD, textColor=colors.white,
                )
                _rl_rows = []
                for ri, row in enumerate(_all_rows):
                    padded = (row + [""] * _num_cols)[:_num_cols]
                    if ri == 0:
                        _rl_rows.append([
                            Paragraph(_md_inline_to_rl(c), _cell_style_hdr) for c in padded
                        ])
                    else:
                        _rl_rows.append([
                            Paragraph(_md_inline_to_rl(c), _cell_style) for c in padded
                        ])
                _avail = A4[0] - 2 * _PAGE_MARGIN
                _col_widths = [_avail / _num_cols] * _num_cols
                _tbl_obj = Table(_rl_rows, colWidths=_col_widths, repeatRows=1)
                _tbl_obj.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
                    ("FONTNAME", (0, 1), (-1, -1), _FONT_BODY),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), _COLOR_PRIMARY),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _COLOR_LIGHT_BG]),
                    ("GRID", (0, 0), (-1, -1), 0.5, _COLOR_BORDER),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(Spacer(1, 2 * mm))
                story.append(_tbl_obj)
                story.append(Spacer(1, 2 * mm))
        else:
            story.append(Paragraph(_md_inline_to_rl(p["content"]), styles["BodyText14"]))

    # ── 6. DISCLAIMER ────────────────────────────────────────────────
    story.append(Spacer(1, 10 * mm))
    story.append(HRFlowable(width="100%", thickness=1, color=_COLOR_BORDER, spaceAfter=4 * mm))
    story.append(Paragraph("6. Disclaimer", styles["SectionHeading"]))
    story.append(Paragraph(
        "This report has been generated using automated data retrieval and AI-powered analysis. "
        "While every effort is made to ensure accuracy, the information contained herein is based on "
        "publicly available data from Companies House and web-based searches. This report does not "
        "constitute legal, financial, or regulatory advice. Recipients should independently verify "
        "key findings before making decisions. AI-generated narrative sections may contain inaccuracies "
        "and should be reviewed by a qualified analyst.",
        styles["Disclaimer"],
    ))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(
        f"Generated by Know Your Company UK on {report_date}. "
        "Built by Ambuj Shukla. Confidential.",
        styles["Disclaimer"],
    ))

    doc.build(story)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════
# DOCX GENERATION — SHARED HELPERS
# ═════════════════════════════════════════════════════════════════════════

def _docx_set_font(run, size=_BODY_SIZE, bold=False, italic=False, color=None):
    """Apply Times New Roman formatting to a DOCX run."""
    run.font.name = _FONT_DOCX
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _docx_add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading with Times New Roman font."""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = _FONT_DOCX
    run.font.size = Pt(_HEADING2_SIZE if level <= 2 else _HEADING3_SIZE)
    run.bold = True
    run.font.color.rgb = _DOCX_PRIMARY
    return h


def _docx_add_para(doc: Document, text: str, bold=False, italic=False,
                    color=None, size=_BODY_SIZE, alignment=None):
    """Add a body paragraph with Times New Roman font."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    # Handle inline bold markers **text**
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _docx_set_font(run, size=size, bold=True, color=color)
        else:
            run = p.add_run(part)
            _docx_set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _docx_add_bullet(doc: Document, text: str, color=None):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _docx_set_font(run, bold=True, color=color)
        else:
            run = p.add_run(part)
            _docx_set_font(run, color=color)
    return p


def _docx_add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add a formatted table to the DOCX document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _docx_set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "1a365d",
        })
        shading.append(shading_elem)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _docx_set_font(run, size=11)

    doc.add_paragraph()  # spacing after table
    return table


def _docx_render_markdown(doc: Document, md_text: str):
    """Render parsed markdown into the DOCX document."""
    paras = _md_to_paragraphs(_clean_md(md_text))
    for p in paras:
        if p["type"] == "h1":
            _docx_add_heading(doc, p["content"], level=1)
        elif p["type"] == "h2":
            _docx_add_heading(doc, p["content"], level=2)
        elif p["type"] == "h3":
            _docx_add_heading(doc, p["content"], level=3)
        elif p["type"] == "bullet":
            _docx_add_bullet(doc, p["content"])
        elif p["type"] == "hr":
            # Add a thin line as paragraph border
            _docx_add_para(doc, "─" * 60, color=_DOCX_GREY, size=8)
        elif p["type"] == "table":
            _tbl_data = p["content"]
            _tbl_headers = _tbl_data.get("headers", [])
            _tbl_rows = _tbl_data.get("rows", [])
            if _tbl_headers:
                _num_cols = len(_tbl_headers)
                # Pad/trim rows to match header count
                _clean_rows = [
                    (row + [""] * _num_cols)[:_num_cols] for row in _tbl_rows
                ]
                _docx_add_table(doc, _tbl_headers, _clean_rows)
        else:
            _docx_add_para(doc, p["content"])


# ═════════════════════════════════════════════════════════════════════════
# DOCX GENERATION — CHARITY
# ═════════════════════════════════════════════════════════════════════════

def generate_charity_docx(dp: Dict[str, Any]) -> bytes:
    """
    Generate a professional DOCX report for a charity.
    Returns: DOCX file content as bytes.
    """
    doc = Document()

    # ── Set default font ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_DOCX
    font.size = Pt(_BODY_SIZE)

    # ── Extract data ─────────────────────────────────────────────────
    entity_name = dp.get("entity_name", "Unknown Charity")
    charity_num = dp.get("charity_num", "")
    charity_data = dp.get("charity_data", {})
    full_report = dp.get("full_report", "")
    trustees = dp.get("trustees", [])
    financial_history = dp.get("financial_history", [])
    financial_anomalies = dp.get("financial_anomalies", {})
    countries = dp.get("countries", [])
    high_risk_countries = dp.get("high_risk_countries", [])
    hrcob_core_controls = dp.get("hrcob_core_controls", {})
    adverse_org = dp.get("adverse_org", {})
    adverse_trustees = dp.get("adverse_trustees", {})
    v3_risk = dp.get("v3_risk_score", {})
    report_date = datetime.now().strftime("%d %B %Y")

    _risk_match = re.search(
        r"(?:Overall\s+Risk\s+Rating)[:\s*#]*\*{0,2}\s*"
        r"(LOW|MEDIUM|HIGH|VERY\s*HIGH|MEDIUM[\s-]*HIGH|MEDIUM[\s-]*LOW)",
        full_report, re.IGNORECASE,
    )
    overall_risk = _risk_match.group(1).strip().upper() if _risk_match else "UNRATED"
    hrcob_stat = hrcob_core_controls.get("hrcob_status", "Unknown")
    inc = charity_data.get("latest_income") or 0
    exp = charity_data.get("latest_expenditure") or 0
    _n_adverse = sum(_count_adverse(v) for v in adverse_trustees.values()) + _count_adverse(adverse_org)

    # ── COVER PAGE ───────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Know Your Charity UK")
    _docx_set_font(run, size=24, bold=True, color=_DOCX_PRIMARY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("In-Depth Due Diligence Report")
    _docx_set_font(run2, size=16, color=_DOCX_ACCENT)

    doc.add_paragraph()  # spacer

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(entity_name)
    _docx_set_font(run3, size=20, bold=True, color=_DOCX_PRIMARY)

    for meta in [
        f"Charity Commission No: {charity_num}",
        f"Report Date: {report_date}",
        f"Registration Status: {charity_data.get('reg_status', 'Unknown')}",
    ]:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_m = pm.add_run(meta)
        _docx_set_font(run_m, size=11, color=_DOCX_GREY)

    doc.add_paragraph()
    pr = doc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = pr.add_run(f"Overall Risk Assessment: {overall_risk}")
    _docx_set_font(run_r, size=16, bold=True, color=_risk_color_docx(overall_risk))

    for _ in range(4):
        doc.add_paragraph()
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run("CONFIDENTIAL — For authorised recipients only")
    _docx_set_font(rc, size=10, italic=True, color=_DOCX_GREY)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────
    _docx_add_heading(doc, "1. Executive Summary", level=1)

    _docx_add_para(doc, f"**Overall Risk Rating:** {overall_risk}",
                    color=_risk_color_docx(overall_risk))
    _docx_add_para(doc, f"**HRCOB Core Controls Status:** {hrcob_stat}")

    _docx_add_table(doc,
        ["Metric", "Value"],
        [
            ["Latest Income", _fmt_money(inc)],
            ["Latest Expenditure", _fmt_money(exp)],
            ["Surplus / (Deficit)", _fmt_money(inc - exp)],
            ["Number of Trustees", str(len(trustees))],
            ["Countries of Operation", str(len(countries))],
            ["High-Risk Jurisdictions", str(len(high_risk_countries))],
            ["Financial Anomalies", str(financial_anomalies.get("anomaly_count", 0) if financial_anomalies else 0)],
        ],
    )

    _docx_add_para(doc, "**Key Findings:**", bold=True)
    _findings = []
    _findings.append(f"HRCOB Core Controls: {hrcob_stat}")
    if high_risk_countries:
        _findings.append(f"{len(high_risk_countries)} high-risk jurisdiction(s): {', '.join(high_risk_countries[:5])}")
    else:
        _findings.append("No high-risk jurisdictions identified")
    if _n_adverse > 0:
        _findings.append(f"{_n_adverse} verified adverse media hit(s) — review required")
    else:
        _findings.append("No verified adverse media hits")
    if inc > 0:
        _spend_pct = exp / inc * 100
        _findings.append(f"Income {_fmt_money(inc)}, Spend-to-Income {_spend_pct:.0f}%")
    for f in _findings:
        _docx_add_bullet(doc, f)

    # ── 2. FINANCIAL OVERVIEW ────────────────────────────────────────
    _docx_add_heading(doc, "2. Financial Overview", level=1)
    _docx_add_heading(doc, "Income Breakdown (Latest Year)", level=2)

    _inc_sources = {
        "Donations & Legacies": charity_data.get("inc_donations", 0) or 0,
        "Charitable Activities": charity_data.get("inc_charitable", 0) or 0,
        "Trading": charity_data.get("inc_trading", 0) or 0,
        "Investments": charity_data.get("inc_investments", 0) or 0,
        "Other": charity_data.get("inc_other", 0) or 0,
    }
    _docx_add_table(doc,
        ["Source", "Amount"],
        [[src, _fmt_money(amt)] for src, amt in _inc_sources.items()] +
        [["Total Income", _fmt_money(inc)]],
    )

    if financial_history:
        _docx_add_heading(doc, "Financial History", level=2)
        _fh_rows = []
        for fh in financial_history[:6]:
            _yr = str(fh.get("year_end", ""))[:10] if fh.get("year_end") else "—"
            _fi = fh.get("income") or 0
            _fe = fh.get("expenditure") or 0
            _fh_rows.append([_yr, _fmt_money(_fi), _fmt_money(_fe), _fmt_money(_fi - _fe)])
        _docx_add_table(doc,
            ["Year End", "Income", "Expenditure", "Surplus/Deficit"],
            _fh_rows,
        )

    _anomaly_ct = financial_anomalies.get("anomaly_count", 0) if financial_anomalies else 0
    if _anomaly_ct > 0:
        _docx_add_para(doc, f"**Financial Anomalies Detected:** {_anomaly_ct}")
        for det in (financial_anomalies.get("details") or [])[:5]:
            _docx_add_bullet(doc, str(det))

    # ── 3. GOVERNANCE & CONTROLS ─────────────────────────────────────
    _docx_add_heading(doc, "3. Governance & Controls", level=1)
    _docx_add_para(doc, f"**HRCOB Core Controls Assessment:** {hrcob_stat}")

    _ctrl_rows = []
    for area in ["safeguarding", "financial_crime", "risk_management"]:
        ctrl = hrcob_core_controls.get(area, {})
        _status = ctrl.get("status", "not_located").replace("_", " ").title()
        _evidence = ctrl.get("evidence", "—")
        if len(_evidence) > 120:
            _evidence = _evidence[:117] + "..."
        _ctrl_rows.append([area.replace("_", " ").title(), _status, _evidence])
    _docx_add_table(doc, ["Control Area", "Status", "Evidence"], _ctrl_rows)

    if trustees:
        _docx_add_heading(doc, "Board of Trustees", level=2)
        _t_rows = []
        for tr in trustees[:15]:
            _tname = tr.get("name", "—") if isinstance(tr, dict) else str(tr)
            _trole = tr.get("role", "Trustee") if isinstance(tr, dict) else "Trustee"
            _t_rows.append([_tname, _trole])
        if len(trustees) > 15:
            _t_rows.append([f"... and {len(trustees) - 15} more", ""])
        _docx_add_table(doc, ["Name", "Role"], _t_rows)

    # ── 4. GEOGRAPHY ─────────────────────────────────────────────────
    _docx_add_heading(doc, "4. Geography Risk Assessment", level=1)
    _country_names = [c.get("country", str(c)) if isinstance(c, dict) else str(c) for c in countries]
    if _country_names:
        _docx_add_para(doc, f"**Countries of Operation:** {', '.join(_country_names[:20])}")
    else:
        _docx_add_para(doc, "No international operations identified.")

    if high_risk_countries:
        _docx_add_para(doc, f"**High-Risk Jurisdictions ({len(high_risk_countries)}):**")
        for c in high_risk_countries:
            _docx_add_bullet(doc, c)
    else:
        _docx_add_para(doc, "No high-risk jurisdictions identified.")

    # ── 5. SCREENING RESULTS ─────────────────────────────────────────
    _docx_add_heading(doc, "5. Screening Results", level=1)
    _docx_add_heading(doc, "Adverse Media Screening", level=2)
    if _n_adverse > 0:
        _docx_add_para(doc,
            f"{_n_adverse} verified adverse media hit(s) detected. Manual review recommended.",
            color=_DOCX_RED)
    else:
        _docx_add_para(doc,
            "No verified adverse media hits detected for the organisation or its trustees.")

    fatf_org = dp.get("fatf_org_screen", {})
    if fatf_org:
        _docx_add_heading(doc, "FATF Screening", level=2)
        _docx_add_para(doc, f"Organisation FATF Result: {fatf_org.get('result', 'No Match')}")

    # ── 6. AI ANALYST REPORT ─────────────────────────────────────────
    doc.add_page_break()
    _docx_add_heading(doc, "6. AI Analyst Report", level=1)
    _docx_add_para(doc,
        "The following analyst report was generated by AI based on structured data, "
        "Charity Commission filings, Companies House records, web searches, and "
        "document analysis.",
        italic=True, size=10, color=_DOCX_GREY)

    _docx_render_markdown(doc, full_report)

    # ── 7. DISCLAIMER ────────────────────────────────────────────────
    _docx_add_heading(doc, "7. Disclaimer", level=1)
    _docx_add_para(doc,
        "This report has been generated using automated data retrieval and AI-powered analysis. "
        "While every effort is made to ensure accuracy, the information contained herein is based on "
        "publicly available data sources including the Charity Commission for England and Wales, "
        "Companies House, FATF lists, and web-based searches. This report does not constitute legal, "
        "financial, or regulatory advice. Recipients should independently verify key findings before "
        "making decisions. The overall risk rating reflects a point-in-time assessment and may change "
        "as new information becomes available. AI-generated narrative sections may contain inaccuracies "
        "and should be reviewed by a qualified analyst.",
        italic=True, size=10, color=_DOCX_GREY)

    _docx_add_para(doc,
        f"Generated by Know Your Charity UK on {report_date}. "
        "Built by Ambuj Shukla. Confidential.",
        italic=True, size=10, color=_DOCX_GREY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════
# DOCX GENERATION — COMPANY
# ═════════════════════════════════════════════════════════════════════════

def generate_company_docx(co_dp: Dict[str, Any]) -> bytes:
    """
    Generate a professional DOCX report for a company.
    Returns: DOCX file content as bytes.
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_DOCX
    font.size = Pt(_BODY_SIZE)

    # ── Extract data ─────────────────────────────────────────────────
    co_data = co_dp.get("co_check_data", {})
    co_report = co_dp.get("co_report", "")
    v3_risk = co_dp.get("v3_risk_score", {})
    report_date = datetime.now().strftime("%d %B %Y")

    co_name = co_data.get("company_name", "Unknown Company")
    co_num = co_data.get("company_number", "")
    co_profile = co_data.get("profile", {})
    co_risk = co_data.get("risk_matrix", {})
    co_age = co_data.get("company_age", {})
    co_directors = co_data.get("director_analysis", {})
    co_pscs = co_data.get("psc_analysis", {})
    co_xref = co_data.get("cross_reference", {})
    co_merchant = co_data.get("merchant_suitability", {})
    co_sic = co_data.get("sic_risk", {})
    co_virtual = co_data.get("virtual_office", {})
    co_ubo = co_data.get("ubo_chain", {})

    _co_overall = co_risk.get("overall_risk", "Unknown")
    _co_flags = co_risk.get("total_flags", 0)
    _co_score = co_risk.get("risk_score", "N/A")
    _co_hard_stops = co_risk.get("hard_stops", [])

    # ── COVER PAGE ───────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Know Your Company UK")
    _docx_set_font(run, size=24, bold=True, color=_DOCX_PRIMARY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Company Sense-Check Report")
    _docx_set_font(run2, size=16, color=_DOCX_ACCENT)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(co_name)
    _docx_set_font(run3, size=20, bold=True, color=_DOCX_PRIMARY)

    for meta in [
        f"Companies House No: {co_num}",
        f"Report Date: {report_date}",
        f"Company Status: {co_profile.get('status', 'Unknown')}",
    ]:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_m = pm.add_run(meta)
        _docx_set_font(run_m, size=11, color=_DOCX_GREY)

    doc.add_paragraph()
    pr = doc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = pr.add_run(f"Overall Risk: {_co_overall.upper()} — Score: {_co_score}/100")
    _docx_set_font(run_r, size=16, bold=True, color=_risk_color_docx(_co_overall))

    for _ in range(4):
        doc.add_paragraph()
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run("CONFIDENTIAL — For authorised recipients only")
    _docx_set_font(rc, size=10, italic=True, color=_DOCX_GREY)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────
    _docx_add_heading(doc, "1. Executive Summary", level=1)
    _docx_add_para(doc, f"**Overall Risk:** {_co_overall.upper()}",
                    color=_risk_color_docx(_co_overall))
    _docx_add_para(doc, f"**Risk Score:** {_co_score}/100  |  **Flags Detected:** {_co_flags}")

    if _co_hard_stops:
        _docx_add_heading(doc, "Hard Stops", level=2)
        for hs in _co_hard_stops:
            _docx_add_bullet(doc, hs, color=_DOCX_RED)

    _docx_add_table(doc,
        ["Metric", "Value"],
        [
            ["Company Status", co_profile.get("status", "—")],
            ["Company Type", co_profile.get("type", "—")],
            ["Company Age", f"{co_age.get('years', '?')} year(s)"],
            ["SIC Codes", ", ".join(co_profile.get("sic_codes", []) or ["—"])],
            ["Director Count", str(co_directors.get("director_count", 0))],
            ["UBO Layers Traced", str(co_ubo.get("layers_traced", 0))],
            ["Business Model", co_merchant.get("business_model", "—")],
            ["Website Credibility", co_xref.get("credibility_level", "Unknown")],
        ],
    )

    # ── 2. COMPANY PROFILE ───────────────────────────────────────────
    _docx_add_heading(doc, "2. Company Profile", level=1)
    _prof_rows = [
        ["Legal Name", co_name],
        ["Company Number", co_num],
        ["Incorporated", co_profile.get("date_of_creation", "—")],
        ["Jurisdiction", co_profile.get("jurisdiction", "—")],
        ["Accounts Next Due", co_profile.get("accounts_next_due", "—")],
        ["Confirmation Next Due", co_profile.get("confirmation_next_due", "—")],
    ]
    _reg = co_profile.get("registered_office", {})
    if _reg:
        _addr = ", ".join(
            _reg.get(k, "") for k in
            ["address_line_1", "address_line_2", "locality", "region", "postal_code"]
            if _reg.get(k)
        )
        _prof_rows.append(["Registered Office", _addr])
    _docx_add_table(doc, ["Field", "Value"], _prof_rows)

    if co_sic:
        _docx_add_para(doc, f"**SIC Risk Level:** {co_sic.get('risk_level', 'Unknown')}")
    if co_virtual:
        _docx_add_para(doc, f"**Virtual Office Indicator:** {co_virtual.get('indicator', 'Unknown')}")

    # ── 3. DIRECTORS & OWNERSHIP ─────────────────────────────────────
    _docx_add_heading(doc, "3. Directors & Ownership", level=1)
    _dirs = co_directors.get("directors", [])
    if _dirs:
        _docx_add_heading(doc, "Active Directors", level=2)
        _d_rows = []
        for d in _dirs[:12]:
            _d_rows.append([
                d.get("name", "—"),
                d.get("officer_role", "—"),
                str(d.get("appointed_on", "—")),
                d.get("nationality", "—"),
            ])
        if len(_dirs) > 12:
            _d_rows.append([f"... and {len(_dirs) - 12} more", "", "", ""])
        _docx_add_table(doc, ["Name", "Role", "Appointed", "Nationality"], _d_rows)

    _pscs = co_pscs.get("pscs", [])
    if _pscs:
        _docx_add_heading(doc, "Persons with Significant Control", level=2)
        _psc_rows = []
        for p in _pscs[:10]:
            _psc_rows.append([
                p.get("name", "—"),
                p.get("nationality", "—"),
                p.get("ownership_band", "—"),
            ])
        _docx_add_table(doc, ["Name", "Nationality", "Ownership"], _psc_rows)

    # ── 4. RISK MATRIX ───────────────────────────────────────────────
    _docx_add_heading(doc, "4. Risk Matrix", level=1)
    _cat_risks = co_risk.get("category_risks", {})
    if _cat_risks:
        _rm_rows = [[cat, level.title()] for cat, level in _cat_risks.items()]
        _docx_add_table(doc, ["Category", "Risk Level"], _rm_rows)

    # ── 5. AI ANALYST REPORT ─────────────────────────────────────────
    doc.add_page_break()
    _docx_add_heading(doc, "5. AI Analyst Report", level=1)
    _docx_add_para(doc,
        "The following analyst report was generated by AI based on Companies House "
        "data, web searches, and cross-referencing analysis.",
        italic=True, size=10, color=_DOCX_GREY)
    _docx_render_markdown(doc, co_report)

    # ── 6. DISCLAIMER ────────────────────────────────────────────────
    _docx_add_heading(doc, "6. Disclaimer", level=1)
    _docx_add_para(doc,
        "This report has been generated using automated data retrieval and AI-powered analysis. "
        "While every effort is made to ensure accuracy, the information contained herein is based on "
        "publicly available data from Companies House and web-based searches. This report does not "
        "constitute legal, financial, or regulatory advice. Recipients should independently verify "
        "key findings before making decisions. AI-generated narrative sections may contain inaccuracies "
        "and should be reviewed by a qualified analyst.",
        italic=True, size=10, color=_DOCX_GREY)

    _docx_add_para(doc,
        f"Generated by Know Your Company UK on {report_date}. "
        "Built by Ambuj Shukla. Confidential.",
        italic=True, size=10, color=_DOCX_GREY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═════════════════════════════════════════════════════════════════════════
# PRIVATE HELPERS
# ═════════════════════════════════════════════════════════════════════════

def _count_adverse(adv: Any) -> int:
    """Count verified adverse media hits (mirrors count_true_adverse in app.py)."""
    if not adv or not isinstance(adv, dict):
        return 0
    results = adv.get("results", [])
    if isinstance(results, list):
        return sum(1 for r in results if isinstance(r, dict) and r.get("is_relevant"))
    return 0
